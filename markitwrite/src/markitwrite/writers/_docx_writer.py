"""Writer for Microsoft Word (.docx) documents."""

from __future__ import annotations

import io
import sys
from typing import BinaryIO, Optional

from markitwrite._base_writer import DocumentWriter, WriteResult

_dependency_exc_info = None
try:
    from docx import Document
    from docx.shared import Inches
except ImportError:
    _dependency_exc_info = sys.exc_info()


class DocxWriter(DocumentWriter):
    """Insert images into Word (.docx) documents.

    Uses python-docx under the hood.

    Position hints:
        {"paragraph": N}  - Insert after paragraph N (0-indexed).
        {"after_text": "..."} - Insert after first paragraph containing text.
        No position = append to end of document.

    Size hints:
        {"width": 6.0}  - Width in inches (height auto-calculated).
        {"width": 4, "height": 3} - Both dimensions in inches.
    """

    ACCEPTED_EXTENSIONS = [".docx"]

    def insert_image(
        self,
        image_stream: BinaryIO,
        document_stream: Optional[BinaryIO] = None,
        target_format: str = ".docx",
        position: Optional[dict] = None,
        size: Optional[dict] = None,
        **kwargs,
    ) -> WriteResult:
        if _dependency_exc_info is not None:
            raise ImportError(
                "python-docx is required for DOCX writing. "
                "Install it with: pip install 'markitwrite[docx]'"
            ) from _dependency_exc_info[1]

        # Load or create document
        if document_stream:
            document_stream.seek(0)
            doc = Document(document_stream)
        else:
            doc = Document()

        # Determine sizing
        width = None
        height = None
        if size:
            if "width" in size:
                width = Inches(float(size["width"]))
            if "height" in size:
                height = Inches(float(size["height"]))
        if width is None and height is None:
            width = Inches(6.0)  # sensible default

        # Determine insertion point
        image_stream.seek(0)
        if position and "paragraph" in position:
            para_idx = int(position["paragraph"])
            paragraphs = doc.paragraphs
            if 0 <= para_idx < len(paragraphs):
                # Insert a new paragraph after the target and add picture there
                target_para = paragraphs[para_idx]
                new_para = _insert_paragraph_after(target_para)
                run = new_para.add_run()
                run.add_picture(image_stream, width=width, height=height)
            else:
                doc.add_picture(image_stream, width=width, height=height)
        elif position and "after_text" in position:
            search_text = position["after_text"]
            inserted = False
            for para in doc.paragraphs:
                if search_text in para.text:
                    new_para = _insert_paragraph_after(para)
                    run = new_para.add_run()
                    run.add_picture(image_stream, width=width, height=height)
                    inserted = True
                    break
            if not inserted:
                doc.add_picture(image_stream, width=width, height=height)
        else:
            doc.add_picture(image_stream, width=width, height=height)

        # Save to bytes
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        return WriteResult(
            output=output.read(),
            target_format=".docx",
            images_inserted=1,
            metadata={"size": size or {"width": 6.0}},
        )


def _insert_paragraph_after(paragraph):
    """Insert a new paragraph element after the given paragraph."""
    from docx.oxml.ns import qn

    new_p = paragraph._element.makeelement(qn("w:p"), {})
    paragraph._element.addnext(new_p)
    # Wrap in a Paragraph object
    from docx.text.paragraph import Paragraph

    return Paragraph(new_p, paragraph._parent)
